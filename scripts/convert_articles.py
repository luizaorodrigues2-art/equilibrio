# -*- coding: utf-8 -*-
"""Convert DOCX articles into structured JSON for the portal."""
from __future__ import annotations

import json
import re
import unicodedata
from datetime import datetime, timedelta
from pathlib import Path

from docx import Document

ARTIGOS_DIR = Path(r"D:\AURYX\08 SITES\equilibrio\artigos")
OUT_DIR = Path(r"D:\AURYX\08 SITES\equilibrio\content\articles")
INDEX_PATH = Path(r"D:\AURYX\08 SITES\equilibrio\content\index.json")
AUTHOR = "Antonio Paulo Tavares Pereira"
SITE = "EQUILÍBRIO INTEGRAL"

CORPO_KEYWORDS = [
    "corpo", "postura", "nutri", "sono", "hidrata", "imunolog", "intestino",
    "movimento", "respiração diafrag", "longevidade saud", "desintoxica",
    "nr-01", "templo físico", "atividade física",
]
MENTE_KEYWORDS = [
    "bem-estar", "bem estar", "autocuidado", "estresse", "saúde mental",
    "saúde para mente", "relações", "longevidade:", "rotina de bem",
]
ESPIRITO_KEYWORDS = [
    "espiritual", "gratidão", "perdão", "santuario", "santuário",
    "detox digital", "pé na terra", "gentileza", "cozinha",
]


def slugify(text: str) -> str:
    text = unicodedata.normalize("NFKD", text)
    text = "".join(c for c in text if not unicodedata.combining(c))
    text = text.lower()
    text = re.sub(r"[^a-z0-9\s-]", "", text)
    text = re.sub(r"[\s_]+", "-", text).strip("-")
    text = re.sub(r"-+", "-", text)
    return text[:80]


def is_meta_para(p: str) -> bool:
    pl = p.lower().strip()
    if len(pl) < 20:
        return False
    meta_phrases = [
        "vamos refazer", "vamos avançar", "vamos fechar", "vamos fazer um por um",
        "google adsense", "guia para criar artigo", "mais de 800 palavras",
        "tom de conversa de café", "dica de ouro, finalizando", "assinatura da série",
        "links internos para", "mantive o mesmo padrão", "mantive exatamente",
        "mantendo o foco", "mantendo rigorosamente", "mantendo o compromisso",
        "com base na sua pauta", "com base nas diretrizes", "com base nas orientações",
        "com base no planejamento", "baseado no documento", "baseado nas diretrizes",
        "inserí sugestões", "inseri sugestões", "como você sugeriu",
        "sem problemas! vamos", "com certeza! vamos",
        "aqui está o artigo", "aqui está o conteúdo completo",
        "o tema de hoje é", "o tema de encerramento é",
    ]
    return any(m in pl for m in meta_phrases)


def clean_paragraphs(paras: list[str]) -> list[str]:
    cleaned: list[str] = []
    started = False
    for p in paras:
        if not started:
            if is_meta_para(p):
                continue
            if re.match(r"^ARTIGO\s*\d+\s*$", p.strip(), re.I):
                continue
            started = True
            cleaned.append(p)
        else:
            if is_meta_para(p) and len(cleaned) < 3:
                continue
            cleaned.append(p)
    return cleaned


def extract_title(paras: list[str], filename: str) -> str:
    for p in paras[:6]:
        if re.match(r"^ARTIGO\s*\d+", p, re.I) and len(p) < 20:
            continue
        if re.match(r"^Introdu", p, re.I):
            continue
        m = re.match(r"^Artigo\s*\d+\s*:\s*(.+)$", p, re.I)
        if m:
            return m.group(1).strip()
        if 15 < len(p) < 160:
            if p.endswith("?") or ":" in p or p[0].isupper():
                return p.strip()
    stem = Path(filename).stem
    stem = re.sub(r"^ARTIGO\s*\d+\s*-?\s*", "", stem, flags=re.I)
    return stem.replace("_", " ").title()


def infer_category(title: str, text: str, filename: str) -> tuple[str, str]:
    blob = (title + " " + text + " " + filename).lower()
    scores = {"corpo": 0, "mente": 0, "espirito": 0}
    for k in CORPO_KEYWORDS:
        if k in blob:
            scores["corpo"] += 2
    for k in MENTE_KEYWORDS:
        if k in blob:
            scores["mente"] += 2
    for k in ESPIRITO_KEYWORDS:
        if k in blob:
            scores["espirito"] += 2

    fn = filename.lower()
    if any(
        x in fn
        for x in [
            "espiritualidade", "5 minutos", "gradi", "gratid", "detox",
            "cozinha", "terra", "dizer", "perd", "gentileza", "santuario",
        ]
    ):
        scores["espirito"] += 10
    if any(
        x in fn
        for x in [
            "bem estar", "autocuidado", "sono afeta", "estresse", "alimenta",
            "atividade", "saude para", "rela", "longividade", "rotina",
        ]
    ):
        scores["mente"] += 8
    if "longevidade saudavel" in fn or "longevidade saudável" in fn:
        scores["corpo"] += 12
        scores["mente"] -= 4
    if any(
        x in fn
        for x in [
            "tempo fisico", "nutri", "ritmo", "movimento", "hidrata",
            "imunolog", "desintoxica", "respira", "intistino", "intestino", "nr-01",
        ]
    ):
        scores["corpo"] += 10
    if "sistema imunolog" in fn:
        scores["corpo"] += 10

    best = max(scores, key=scores.get)
    mapping = {
        "corpo": ("Saúde do Corpo", "saude-do-corpo"),
        "mente": ("Saúde da Mente", "saude-da-mente"),
        "espirito": ("Saúde Espiritual", "saude-espiritual"),
    }
    return mapping[best]


def extract_tags(title: str, text: str, category_slug: str) -> list[str]:
    pool = {
        "saude-do-corpo": ["corpo", "saude-fisica", "habitos", "energia"],
        "saude-da-mente": ["mente", "bem-estar", "equilibrio-emocional", "saude-mental"],
        "saude-espiritual": ["espiritualidade", "paz-interior", "autoconhecimento", "presenca"],
    }
    tags = list(pool.get(category_slug, ["bem-estar"]))
    keyword_map = {
        "ansiedade": "ansiedade",
        "estresse": "estresse",
        "sono": "sono",
        "medita": "meditacao",
        "respir": "respiracao",
        "aliment": "alimentacao",
        "nutri": "nutricao",
        "exerc": "exercicio",
        "postura": "postura",
        "perdão": "perdao",
        "perdao": "perdao",
        "gratid": "gratidao",
        "digital": "detox-digital",
        "natureza": "natureza",
        "imun": "imunidade",
        "longevid": "longevidade",
        "hidrata": "hidratacao",
        "rotina": "rotina",
        "autocuidado": "autocuidado",
        "comunica": "comunicacao",
        "relaç": "relacionamentos",
        "relac": "relacionamentos",
        "intestino": "intestino-cerebro",
        "desintox": "detox",
        "gentileza": "gentileza",
        "espiritual": "espiritualidade",
        "paz": "paz",
    }
    blob = (title + " " + text[:1500]).lower()
    for k, tag in keyword_map.items():
        if k in blob and tag not in tags:
            tags.append(tag)
    return tags[:8]


def reading_time(text: str) -> int:
    words = len(text.split())
    return max(1, round(words / 200))


def build_excerpt(paras: list[str], title: str) -> str:
    for p in paras:
        if p == title:
            continue
        if re.match(r"^introdu", p, re.I):
            continue
        if len(p) > 80:
            return (p[:220] + "…") if len(p) > 220 else p
    return title


def to_html_sections(paras: list[str], title: str):
    html_parts: list[str] = []
    sections: list[dict] = []
    intro_paras: list[str] = []
    body_started = False
    skip_author = re.compile(r"^autor\s*:", re.I)

    for i, p in enumerate(paras):
        if p.strip() == title:
            continue
        if skip_author.match(p):
            continue
        if re.match(r"^ARTIGO\s*\d+\s*$", p, re.I):
            continue

        is_h2 = False
        is_h3 = False

        if re.match(r"^(introdu[cç][aã]o|conclus[aã]o|considera[cç][oõ]es finais)", p, re.I) and len(p) < 80:
            is_h2 = True
        elif re.match(r"^dica de ouro", p, re.I):
            is_h2 = True
        elif re.match(r"^\d{1,2}[\.\)]\s+", p) and len(p) < 100:
            is_h2 = True
        elif len(p) < 90 and not p.endswith(".") and not p.endswith(",") and i > 0:
            if i + 1 < len(paras) and len(paras[i + 1]) > 100:
                if p[0].isupper() and (
                    ":" in p
                    or p.istitle()
                    or sum(1 for c in p if c.isupper()) > 2
                    or len(p.split()) <= 12
                ):
                    is_h2 = True
        elif len(p) < 70 and p.endswith(":") and i > 0:
            is_h3 = True

        if is_h2:
            body_started = True
            sid = slugify(p)
            sections.append({"id": sid, "title": p})
            html_parts.append(f'<h2 id="{sid}">{p}</h2>')
        elif is_h3:
            body_started = True
            sid = slugify(p)
            html_parts.append(f'<h3 id="{sid}">{p}</h3>')
        else:
            if html_parts and "dica-de-ouro" in html_parts[-1].lower():
                html_parts.append(f'<div class="callout callout--gold"><p>{p}</p></div>')
            elif not body_started and not intro_paras and len(p) > 100:
                intro_paras.append(p)
                html_parts.append(
                    f'<p class="lead"><span class="dropcap">{p[0]}</span>{p[1:]}</p>'
                )
            else:
                if not body_started:
                    intro_paras.append(p)
                html_parts.append(f"<p>{p}</p>")

    return "\n".join(html_parts), sections, intro_paras


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    articles: list[dict] = []
    base_date = datetime(2026, 1, 15)

    for idx, f in enumerate(sorted(ARTIGOS_DIR.glob("*.docx"))):
        doc = Document(str(f))
        paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paras.append(" | ".join(cells))

        paras = clean_paragraphs(paras)
        title = extract_title(paras, f.name)
        title = re.sub(r"^Artigo\s*\d+\s*:\s*", "", title, flags=re.I).strip()
        full_text = "\n\n".join(paras)
        cat_name, cat_slug = infer_category(title, full_text, f.name)
        tags = extract_tags(title, full_text, cat_slug)
        html, toc, _intro = to_html_sections(paras, title)
        excerpt = build_excerpt(paras, title)

        subtitle = ""
        for p in paras:
            if p != title and len(p) > 40:
                m = re.match(r"^(.{40,140}[\.!?])", p)
                subtitle = m.group(1) if m else p[:140]
                break

        slug = slugify(title)
        existing = {a["slug"] for a in articles}
        base_slug = slug
        n = 2
        while slug in existing:
            slug = f"{base_slug}-{n}"
            n += 1

        pub = base_date + timedelta(days=idx * 3)
        cover = f"/images/covers/{cat_slug}.jpg"

        article = {
            "id": f"art-{idx + 1:03d}",
            "slug": slug,
            "title": title,
            "subtitle": subtitle,
            "excerpt": excerpt,
            "content": html,
            "contentText": full_text,
            "author": AUTHOR,
            "category": cat_name,
            "categorySlug": cat_slug,
            "tags": tags,
            "coverImage": cover,
            "coverAlt": f"Ilustração do artigo: {title}",
            "publishedAt": pub.strftime("%Y-%m-%d"),
            "updatedAt": pub.strftime("%Y-%m-%d"),
            "readingTime": reading_time(full_text),
            "featured": idx < 3,
            "status": "published",
            "views": 1000 - idx * 23,
            "likes": 40 + (idx % 17) * 3,
            "seo": {
                "title": f"{title} | {SITE}",
                "description": excerpt[:160],
                "keywords": tags + [cat_name.lower(), "equilibrio integral", "bem-estar"],
            },
            "toc": toc,
            "sourceFile": f.name,
            "faq": [],
        }
        articles.append(article)
        print(f"[{cat_slug}] {slug} | {article['readingTime']}min | {len(full_text)} chars")
        (OUT_DIR / f"{slug}.json").write_text(
            json.dumps(article, ensure_ascii=False, indent=2), encoding="utf-8"
        )

    index = [
        {
            k: a[k]
            for k in [
                "id", "slug", "title", "subtitle", "excerpt", "author", "category",
                "categorySlug", "tags", "coverImage", "coverAlt", "publishedAt",
                "readingTime", "featured", "status", "views", "likes", "seo",
            ]
        }
        for a in articles
    ]
    INDEX_PATH.parent.mkdir(parents=True, exist_ok=True)
    INDEX_PATH.write_text(json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8")

    cats: dict[str, int] = {}
    for a in articles:
        cats[a["category"]] = cats.get(a["category"], 0) + 1
    print(f"\nSaved {len(articles)} articles")
    print("Categories:", cats)


if __name__ == "__main__":
    main()
